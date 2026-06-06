"""Generate CLIENT_USER_AND_ROLES_GUIDE.docx — run once locally."""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

OUT = Path(__file__).resolve().parent / "CLIENT_USER_AND_ROLES_GUIDE.docx"

NAVY = RGBColor(0x0F, 0x27, 0x44)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = NAVY
    return h


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, text in enumerate(headers):
        hdr[i].text = text
        for p in hdr[i].paragraphs:
            for run in p.runs:
                run.bold = True
    for r_idx, row in enumerate(rows):
        cells = table.rows[r_idx + 1].cells
        for c_idx, text in enumerate(row):
            cells[c_idx].text = str(text)
    doc.add_paragraph()


def main():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    title = doc.add_heading("Sama Tours ERP", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = NAVY

    sub = doc.add_paragraph(
        "Guide for owners & managers — Users, Groups, Employees, and access\n"
        "Sama Accounting system"
    )
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in sub.runs:
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    doc.add_paragraph()

    add_heading(doc, "What is a “Group”? Do I need to use it?", 1)
    doc.add_paragraph(
        "Yes — you should use Groups. A Group is simply a job role attached to a login account. "
        "The system has three groups:"
    )
    for item in [
        "Sales — sales staff",
        "Accounting — finance / accountants",
        "Admin — manager or owner with full control",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_paragraph(
        "When you create a user, you pick one Group. That decides what they can see in the ERP "
        "(for example, Sales cannot see supplier costs on invoices)."
    )
    p = doc.add_paragraph()
    run = p.add_run("Groups are required if you want different access for sales vs accounting.")
    run.bold = True

    doc.add_paragraph(
        "In the Admin panel, Groups appear under Authentication and Authorization → Groups. "
        "You normally do not create new groups — only assign users to Sales, Accounting, or Admin."
    )

    add_heading(doc, "Two different things (do not confuse them)", 1)
    add_table(
        doc,
        ["", "Login User + Group", "Employee (staff list)"],
        [
            ("Where", "Admin → Users → Groups", "Admin → Employees"),
            (
                "Purpose",
                "Username, password, and permissions",
                "Name on invoices and salesman reports",
            ),
            (
                "Example",
                "“Rana” logs in as Sales",
                "“Rana” appears as salesperson on an invoice",
            ),
        ],
    )
    doc.add_paragraph(
        "The Employee role (Sales / Accounting / Admin on the employee record) is for reporting labels, "
        "not for login security. Login security comes from the user’s Group."
    )

    add_heading(doc, "Two screens in the system", 1)
    add_table(
        doc,
        ["Screen", "Address", "Who uses it"],
        [
            ("Main ERP (daily work)", "/login/ — Dashboard, Invoices, etc.", "Everyone with a user account"),
            ("Admin panel (settings)", "/admin/", "Only Staff users (owner, IT, lead accountant)"),
        ],
    )

    add_heading(doc, "What each Group can do in the ERP", 1)
    add_table(
        doc,
        ["Area", "Sales", "Accounting", "Admin"],
        [
            ("Dashboard & reports", "Yes", "Yes", "Yes"),
            ("Clients", "Yes", "Yes", "Yes"),
            ("Create / edit draft invoices", "Yes", "Yes", "Yes"),
            ("See supplier cost & profit", "No", "Yes", "Yes"),
            ("Accountant invoice PDF", "No", "Yes", "Yes"),
            ("Bills, payments, expenses", "Not for daily use", "Yes", "Yes"),
            ("Adjust posted invoices", "No", "Yes", "Yes"),
            ("Full /admin/ panel", "No", "If Staff ticked", "If Staff ticked"),
        ],
    )

    add_heading(doc, "Who should access the Admin panel (/admin/)?", 1)
    add_table(
        doc,
        ["Person", "Group", "Staff?", "Superuser?"],
        [
            ("Developer / IT", "—", "Yes", "Yes (you only)"),
            ("CEO / owner", "Admin", "Yes", "No"),
            ("Main accountant", "Accounting", "Optional", "No"),
            ("Other accountants", "Accounting", "No", "No"),
            ("Sales employees", "Sales", "No", "No"),
        ],
    )
    doc.add_paragraph(
        "Staff = checkbox on the user page; required to open /admin/.\n"
        "Superuser = full technical control; keep for 1–2 people only."
    )

    add_heading(doc, "“Is main accountant” — should I use it?", 1)
    doc.add_paragraph(
        "Yes — for one person only: your lead / head accountant."
    )
    doc.add_paragraph(
        "When adding or editing a user in Admin, under Profile, tick Is main accountant for that person only."
    )
    doc.add_paragraph(
        "Effect: New invoices automatically default the salesperson/employee field to that person’s "
        "Employee record (if linked to their user account)."
    )
    doc.add_paragraph("Also for the main accountant:")
    for step in [
        "Admin → Employees → create or edit their record (role Accounting).",
        "Link User = their login account.",
        "User → Group = Accounting.",
        "Tick Is main accountant on their profile.",
    ]:
        doc.add_paragraph(step, style="List Number")
    doc.add_paragraph("Other accountants: Group Accounting, do not tick main accountant.")

    add_heading(doc, "Step-by-step: new sales employee", 1)
    for step in [
        "Admin → Employees → Add → name, role Sales, active.",
        "Admin → Users → Add user → username & password.",
        "Staff = OFF. Superuser = OFF.",
        "Groups → choose Sales only → Save.",
        "Edit Employee → link User to this account.",
        "Send them the main site URL and login (not /admin/).",
    ]:
        doc.add_paragraph(step, style="List Number")

    add_heading(doc, "Step-by-step: main accountant", 1)
    for step in [
        "Admin → Users → Add user → username & password.",
        "Group = Accounting. Staff = optional. Superuser = OFF.",
        "Tick Is main accountant → Save.",
        "Admin → Employees → add name, role Accounting, link User.",
    ]:
        doc.add_paragraph(step, style="List Number")

    add_heading(doc, "Step-by-step: CEO / owner", 1)
    for step in [
        "User with Group Admin and Staff ticked.",
        "Uses main ERP and /admin/ to manage users if needed.",
        "Do not give every employee superuser access.",
    ]:
        doc.add_paragraph(step, style="List Number")

    add_heading(doc, "Quick reference", 1)
    add_table(
        doc,
        ["Question", "Answer"],
        [
            ("What is a Group?", "A role: Sales, Accounting, or Admin — assigned to each user."),
            ("Must I use Groups?", "Yes, for every login user."),
            ("Where do I assign it?", "Admin → Users → [user] → Groups."),
            ("Employee vs Group?", "Employee = name on documents; Group = permissions."),
            ("Main accountant checkbox?", "One person only; defaults employee on new invoices."),
        ],
    )

    doc.add_paragraph()
    foot = doc.add_paragraph("Sama Tours — Sama Accounting ERP · User & roles guide · For internal use")
    foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in foot.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    doc.save(OUT)
    print(f"Created: {OUT}")


if __name__ == "__main__":
    main()
